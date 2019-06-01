import requests
import os
from bs4 import BeautifulSoup
import urllib
from docx import Document
from docx.shared import Inches

from wang.yu.common.PyXmlParser import PyXmlParser


class Test51:
    def __init__(self):
        py_xml_parser = PyXmlParser(os.path.dirname(os.path.dirname(__file__)) + '/common/properties/properties.xml')
        self.xml = py_xml_parser.do_parse()
        self.images = []
        self.images_location = os.path.dirname(__file__) + '/images'

    def create_word_doc(self):
        document = Document()
        document.add_heading('2019中考压轴100题(上)', 0)

        max_count = self.images.__len__() + 1
        for count in range(1, max_count):
            document.add_picture('%s/%s.png' % (self.images_location, count), width=Inches(5))

        document.save(os.path.dirname(__file__) + '/demo.docx')

    def crawl_images(self):
        test51_yazhou_zhongkao2019_url = self.xml['websites']['test51']['yazhou']['zhongkao2019']['url']
        page_html = requests.get(test51_yazhou_zhongkao2019_url)
        page_html_text = page_html.text
        soup = BeautifulSoup(page_html_text, 'html.parser')
        count = 1
        for img_div in soup.select('p img'):
            img_src = img_div['data-src']
            img_name = '%s.png' % count
            self.download_images(img_src, img_name)
            self.images.append(img_name)
            count += 1
        self.create_word_doc()

    def download_images(self, img_src, img_name):
        if not os.path.exists(self.images_location):
            os.mkdir(self.images_location)
            print('正在创建本地文件夹...')
        if not os.path.exists('%s/%s' % (self.images_location, img_name)):
            try:
                urllib.request.urlretrieve(img_src, '%s/%s' % (self.images_location, img_name))
            except ValueError:
                urllib.request.urlretrieve('%s%s' % ('https:', img_src), '%s/%s' % (self.images_location, img_name))
        else:
            print('已经存在%s' % img_name)


test = Test51()
test.crawl_images()
